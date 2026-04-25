from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "SleepWatch_Dokumentacja_ver_1.md"
TARGET_PATH = BASE_DIR / "SleepWatch_Dokumentacja_ver_1.docx"


def add_image(document: Document, line: str) -> bool:
    stripped = line.strip()
    if not (stripped.startswith("![") and "](" in stripped and stripped.endswith(")")):
        return False

    target = stripped.split("](", 1)[1][:-1].strip()
    image_path = (SOURCE_PATH.parent / target).resolve()
    if not image_path.exists():
        return False

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    return True


def add_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if cells and not all(set(cell) <= {"-"} for cell in cells):
            rows.append(cells)

    if not rows:
        return

    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = value

    for row_values in rows[1:]:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = value


def build_docx() -> None:
    document = Document()
    lines = SOURCE_PATH.read_text(encoding="utf-8").splitlines()

    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[str] = []

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                for code_line in code_buffer:
                    paragraph.add_run(code_line + "\n").font.name = "Consolas"
                code_buffer = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(stripped)
            continue

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue

        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

        if add_image(document, stripped):
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=3)
            continue

        if stripped[:3].isdigit() and stripped[1:3] == ". ":
            document.add_paragraph(stripped, style="List Number")
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        document.add_paragraph(stripped)

    if table_buffer:
        add_table(document, table_buffer)

    document.save(TARGET_PATH)


if __name__ == "__main__":
    build_docx()
